"""DOCX rendering, via python-docx.

The same `Layout` the PDF renderer walks, so the two files say the same thing.
Where they differ is on purpose: a DOCX is an *editable* deliverable. Somebody
receiving one will paste it into a longer incident write-up, so this renderer
leans on Word's built-in styles (`Title`, `Heading 1`, `Heading 2`) rather than
hard-coding every run. That way the document inherits the recipient's template
and its headings show up in Word's navigation pane, instead of arriving as
fifty paragraphs of manually-sized text that fight whatever they are pasted
into.

The n8n Simple Report Generator builds a DOCX too. It reads the prototype's
shape — `parsed_report.sections.executive_summary`, one attack, a confidence
score this schema does not have — and it renders a fixed logo from a path
inside the n8n container. It is kept as reference, not reused.
"""

from __future__ import annotations

import io

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.export.layout import SEVERITY_LABELS, Layout, Section

INK = RGBColor(0x15, 0x18, 0x1E)
DIM = RGBColor(0x59, 0x61, 0x6F)
FAINT = RGBColor(0x82, 0x8B, 0x99)

SEVERITY_COLORS = {
    "critical": RGBColor(0xB0, 0x1F, 0x36),
    "high": RGBColor(0xA1, 0x52, 0x0D),
    "medium": RGBColor(0x83, 0x64, 0x05),
    "low": RGBColor(0x0F, 0x6C, 0x7B),
    "unknown": FAINT,
}

CLASSIFICATION_COLORS = {
    "Confidential": RGBColor(0xB0, 0x1F, 0x36),
    "Internal": INK,
    "Public": DIM,
}

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"


def render(layout: Layout) -> bytes:
    document = Document()
    _configure(document, layout)

    _title_block(document, layout)
    _meta_table(document, layout)
    _summary(document, layout)

    for index, section in enumerate(layout.sections):
        if index == 0:
            document.add_section(WD_SECTION.NEW_PAGE)
            _page_furniture(document.sections[-1], layout.classification)
        _section(document, section)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _configure(document: Document, layout: Layout) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)

    document.core_properties.title = layout.title
    document.core_properties.author = "AIPCC"
    document.core_properties.category = layout.classification
    # Word surfaces "Comments" in the file's properties pane, which is where
    # somebody checks a document's handling caveat before they forward it.
    document.core_properties.comments = f"Classification: {layout.classification}"

    _page_furniture(document.sections[0], layout.classification)


def _page_furniture(section, classification: str) -> None:
    """Classification in the header, classification and page number in the footer.

    Repeated on every page rather than stated once at the top: a document is
    read one page at a time and forwarded one page at a time, and a caveat that
    only appears on page one stops travelling with the pages that leave it.
    """
    colour = CLASSIFICATION_COLORS.get(classification, INK)
    caveat = classification.upper()

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(caveat)
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = colour

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caveat_run = footer.add_run(f"{caveat}    ·    Page ")
    caveat_run.bold = True
    caveat_run.font.size = Pt(8)
    caveat_run.font.color.rgb = colour
    _page_number_field(footer)
    of_run = footer.add_run(" of ")
    of_run.font.size = Pt(8)
    of_run.font.color.rgb = FAINT
    _page_number_field(footer, instruction="NUMPAGES")


def _page_number_field(paragraph, instruction: str = "PAGE") -> None:
    """Insert a Word field, so the number is computed when the file is opened.

    python-docx has no API for this; the three-element `fldChar` dance below is
    what Word itself writes. Rendering a literal number instead would be wrong
    the moment anyone edited the document, which is the whole point of handing
    someone a DOCX.
    """
    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = FAINT

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set(qn("xml:space"), "preserve")
    instruction_element.text = f" {instruction} "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instruction_element)
    run._r.append(end)


# --- Content --------------------------------------------------------------


def _title_block(document: Document, layout: Layout) -> None:
    heading = document.add_heading(layout.title, level=0)
    for run in heading.runs:
        run.font.color.rgb = INK

    if layout.provenance:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(layout.provenance)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = DIM


def _meta_table(document: Document, layout: Layout) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in layout.meta:
        cells = table.add_row().cells
        label_run = cells[0].paragraphs[0].add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(8.5)
        label_run.font.color.rgb = DIM
        value_run = cells[1].paragraphs[0].add_run(value)
        value_run.font.size = Pt(9)
        if label in {"Reference", "Sealed SHA-256"}:
            value_run.font.name = MONO_FONT
            value_run.font.size = Pt(8)


def _summary(document: Document, layout: Layout) -> None:
    document.add_paragraph()
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    header = table.rows[0].cells
    for index, text in enumerate(("Findings by severity", "", "Contents", "")):
        run = header[index].paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = DIM

    rows = max(len(layout.tally), len(layout.totals))
    for index in range(rows):
        cells = table.add_row().cells
        if index < len(layout.tally):
            label, count = layout.tally[index]
            key = next(k for k, v in SEVERITY_LABELS.items() if v == label)
            run = cells[0].paragraphs[0].add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = SEVERITY_COLORS[key]
            _number(cells[1], count)
        if index < len(layout.totals):
            label, count = layout.totals[index]
            run = cells[2].paragraphs[0].add_run(label)
            run.font.size = Pt(9)
            run.font.color.rgb = DIM
            _number(cells[3], count)


def _number(cell, value: int) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(str(value))
    run.font.size = Pt(9)


def _section(document: Document, section: Section) -> None:
    document.add_heading(section.title, level=1)

    caption = document.add_paragraph()
    caption_run = caption.add_run(section.caption)
    caption_run.italic = True
    caption_run.font.size = Pt(8.5)
    caption_run.font.color.rgb = DIM

    if section.count == 0:
        note = document.add_paragraph()
        note_run = note.add_run(section.empty_note)
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = FAINT
        return

    if section.is_table:
        _table(document, section)
        return

    for finding in section.findings:
        heading = document.add_heading(level=2)
        heading_run = heading.add_run(finding.heading)
        heading_run.font.color.rgb = INK
        if finding.severity:
            severity_run = heading.add_run(
                f"   {SEVERITY_LABELS[finding.severity].upper()}"
            )
            severity_run.font.size = Pt(8)
            severity_run.bold = True
            severity_run.font.color.rgb = SEVERITY_COLORS[finding.severity]

        if finding.subtitle:
            subtitle = document.add_paragraph()
            subtitle_run = subtitle.add_run(finding.subtitle)
            subtitle_run.font.name = MONO_FONT
            subtitle_run.font.size = Pt(8.5)
            subtitle_run.font.color.rgb = DIM

        for label, value in finding.fields:
            paragraph = document.add_paragraph()
            label_run = paragraph.add_run(f"{label}  ")
            label_run.bold = True
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = FAINT
            value_run = paragraph.add_run(value)
            value_run.font.size = Pt(9.5)


def _table(document: Document, section: Section) -> None:
    table = document.add_table(rows=1, cols=len(section.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for index, column in enumerate(section.columns):
        run = table.rows[0].cells[index].paragraphs[0].add_run(column.upper())
        run.bold = True
        run.font.size = Pt(7.5)
        run.font.color.rgb = DIM

    _repeat_header(table.rows[0])

    for row in section.rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            run = cells[index].paragraphs[0].add_run(value)
            run.font.size = Pt(8.5)


def _repeat_header(row) -> None:
    """Mark the header row so Word repeats it across a page break.

    A findings table that runs to a second page with no column headings is a
    table an analyst has to scroll back to read, which is the difference
    between a document and a data dump.
    """
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)
